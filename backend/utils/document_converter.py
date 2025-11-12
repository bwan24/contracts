import os
import re
from typing import Optional
import docx
import pdfplumber
from markdownify import markdownify
from fastapi import HTTPException

class DocumentConverter:
    """文档转换器，支持Word和PDF转换为Markdown"""
    
    @staticmethod
    def convert_docx_to_markdown(file_path: str) -> str:
        """将Word文档转换为Markdown格式"""
        try:
            # 打开Word文档
            doc = docx.Document(file_path)
            
            # 存储Markdown内容
            markdown_content = []
            
            # 遍历文档中的每个段落
            for paragraph in doc.paragraphs:
                # 获取段落文本
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # 根据段落样式判断标题级别
                if paragraph.style.name.startswith('Heading'):
                    # 提取标题级别（1-9）
                    heading_level_match = re.match(r'Heading (\d+)', paragraph.style.name)
                    if heading_level_match:
                        level = int(heading_level_match.group(1))
                        if 1 <= level <= 9:
                            markdown_content.append(f"{'#' * level} {text}")
                            continue
                
                # 检查是否是列表项
                is_bullet = False
                for run in paragraph.runs:
                    if run.font.bold:
                        text = f"**{text}**"
                    if run.font.italic:
                        text = f"*{text}*"
                
                # 添加普通段落
                markdown_content.append(text)
            
            # 遍历文档中的每个表格
            for table in doc.tables:
                # 开始表格
                table_md = []
                
                # 添加表头
                header_cells = table.rows[0].cells
                table_md.append(f"| {' | '.join(cell.text.strip() for cell in header_cells)} |")
                table_md.append(f"| {' | '.join(['---'] * len(header_cells))} |")
                
                # 添加表格内容
                for row in table.rows[1:]:
                    table_md.append(f"| {' | '.join(cell.text.strip() for cell in row.cells)} |")
                
                # 将表格添加到Markdown内容中
                markdown_content.extend(table_md)
            
            # 合并所有内容，用空行分隔
            return '\n\n'.join(markdown_content)
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to convert Word document to Markdown: {str(e)}"
            )
    
    @staticmethod
    def convert_pdf_to_markdown(file_path: str) -> str:
        """将PDF文档转换为Markdown格式"""
        try:
            # 存储Markdown内容
            markdown_content = []
            
            with pdfplumber.open(file_path) as pdf:
                # 遍历每一页
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取文本
                    text = page.extract_text()
                    
                    if text:
                        # 清理文本
                        cleaned_text = DocumentConverter._clean_pdf_text(text)
                        
                        # 尝试识别标题
                        lines = cleaned_text.split('\n')
                        for i, line in enumerate(lines):
                            line = line.strip()
                            if not line:
                                continue
                                
                            # 简单的标题识别：较长的行且后面跟着空行可能是标题
                            is_heading = False
                            if len(line) < 100:  # 标题通常不会太长
                                # 检查是否有合适的长度和位置特征
                                if i < len(lines) - 1 and not lines[i + 1].strip():
                                    # 判断可能的标题级别
                                    if i == 0 and page_num == 1:  # 第一页第一行可能是文档标题
                                        markdown_content.append(f"# {line}")
                                        is_heading = True
                                    elif len(line) < 50:  # 较短的行可能是子标题
                                        markdown_content.append(f"## {line}")
                                        is_heading = True
                                
                            # 如果不是标题，添加为普通文本
                            if not is_heading:
                                markdown_content.append(line)
                        
                        # 如果是多页文档，添加分页标记
                        if page_num < len(pdf.pages):
                            markdown_content.append("\n---\n")
            
            # 合并所有内容
            return '\n'.join(markdown_content)
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to convert PDF document to Markdown: {str(e)}"
            )
    
    @staticmethod
    def _clean_pdf_text(text: str) -> str:
        """清理从PDF提取的文本"""
        # 移除多余的空行
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        # 移除每行开头和结尾的空白字符
        lines = [line.strip() for line in text.split('\n')]
        
        # 合并被错误分割的段落（简单逻辑：如果行不以句号、问号或感叹号结尾，且下一行不是空的，则合并）
        merged_lines = []
        current_line = ""
        
        for i, line in enumerate(lines):
            if not line:
                if current_line:
                    merged_lines.append(current_line)
                    current_line = ""
                merged_lines.append("")
                continue
                
            # 如果当前行不为空，且不以句号、问号或感叹号结尾，尝试合并
            if current_line and not current_line.strip().endswith(('.', '?', '!', '。', '？', '！')):
                # 检查是否是列表项或其他不应该合并的情况
                if not (line.startswith(('-', '*', '•', '1.', '2.', '3.', '一、', '二、', '三、'))):
                    current_line += ' ' + line
                    continue
            
            if current_line:
                merged_lines.append(current_line)
            current_line = line
        
        if current_line:
            merged_lines.append(current_line)
        
        return '\n'.join(merged_lines)
    
    @staticmethod
    def convert_file_to_markdown(file_path: str) -> str:
        """根据文件扩展名转换为Markdown"""
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        # 根据扩展名调用相应的转换方法
        if ext == '.docx':
            return DocumentConverter.convert_docx_to_markdown(file_path)
        elif ext == '.pdf':
            return DocumentConverter.convert_pdf_to_markdown(file_path)
        elif ext == '.txt':
            # 直接读取文本文件
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # 尝试其他编码
                try:
                    with open(file_path, 'r', encoding='gbk') as f:
                        return f.read()
                except Exception as e:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to read text file: {str(e)}"
                    )
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type for conversion: {ext}"
            )

# 创建转换器实例
converter = DocumentConverter()